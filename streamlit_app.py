import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# Function to add text with formatting to Word
def add_formatted_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    highlight = run.font._element
    highlight_props = OxmlElement("w:highlight")
    highlight_props.set(qn("w:val"), "yellow")
    highlight.append(highlight_props)

# Function to check if text is bolded and highlighted
def is_bold_and_highlighted(run):
    bold = run.bold
    highlight = run.font._element.find(qn("w:highlight"))
    return bold and highlight is not None

# Function to process input text and generate a Word document
def process_text(input_text):
    doc = Document()
    formatted_text = []
    for line in input_text.strip().split("\n"):
        if line.startswith("*"):
            formatted_text.append(f"**{line[1:].strip()}** (highlighted)")
            add_formatted_paragraph(doc, line[1:].strip())  # Remove the asterisk
        else:
            formatted_text.append(line)
            doc.add_paragraph(line)
    return doc, formatted_text

# Function to process a Word document and add asterisks to bolded/highlighted text
def process_uploaded_word(doc):
    new_doc = Document()
    formatted_text = []
    for paragraph in doc.paragraphs:
        new_paragraph_text = ""
        for run in paragraph.runs:
            if is_bold_and_highlighted(run):
                formatted_text.append(f"*{run.text}* (highlighted and bolded)")
                new_paragraph_text += f"*{run.text}"
            else:
                new_paragraph_text += run.text
                formatted_text.append(run.text)
        new_doc.add_paragraph(new_paragraph_text)
    return new_doc, formatted_text

# Streamlit app
st.title("Text Formatter and Word File Generator")
st.write("Upload a Word or text file. Lines starting with an asterisk (*) will be bolded and highlighted in yellow. Bolded and highlighted text will have an asterisk added before it.")

# File upload
uploaded_file = st.file_uploader("Upload your file", type=["txt", "docx"])

if uploaded_file:
    if uploaded_file.name.endswith(".txt"):
        # Process text file
        input_text = uploaded_file.read().decode("utf-8")
        formatted_doc, preview_text = process_text(input_text)
    elif uploaded_file.name.endswith(".docx"):
        # Process Word file
        uploaded_doc = Document(uploaded_file)
        formatted_doc, preview_text = process_uploaded_word(uploaded_doc)
    else:
        st.warning("Unsupported file format!")
        formatted_doc, preview_text = None, None

    if formatted_doc:
        # Preview the processed text
        st.subheader("Preview of Formatted Text:")
        for line in preview_text:
            st.write(line)

        # Save the Word file in memory
        output = io.BytesIO()
        formatted_doc.save(output)
        output.seek(0)
        # Provide a download link
        st.download_button(label="Download Formatted Word File", data=output, file_name="formatted_text.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Please upload a file to process.")
